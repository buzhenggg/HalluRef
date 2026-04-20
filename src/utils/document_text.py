"""Extract reviewable plain text from uploaded PDF and DOCX files."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Callable

from docx import Document
from pypdf import PdfReader


class DocumentTextError(ValueError):
    """Base error for document text extraction."""


class UnsupportedDocumentError(DocumentTextError):
    """Raised when a file type is not supported."""


class EmptyDocumentTextError(DocumentTextError):
    """Raised when no usable text can be extracted."""


class DocumentParseError(DocumentTextError):
    """Raised when a supported document cannot be parsed."""


_DOCX_MIMES = {
    "",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
}
_PDF_MIMES = {"", "application/pdf", "application/octet-stream"}


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    normalized = "\n".join(lines).strip()
    return re.sub(r"\n{3,}", "\n\n", normalized)


def _validate_type(filename: str, content_type: str | None) -> str:
    suffix = Path(filename or "").suffix.lower()
    mime = (content_type or "").split(";")[0].strip().lower()

    if suffix == ".pdf":
        if mime not in _PDF_MIMES:
            raise UnsupportedDocumentError(f"Unsupported PDF content type: {content_type}")
        return "pdf"

    if suffix == ".docx":
        if mime not in _DOCX_MIMES:
            raise UnsupportedDocumentError(f"Unsupported DOCX content type: {content_type}")
        return "docx"

    if suffix == ".doc":
        raise UnsupportedDocumentError("Unsupported legacy Word .doc files. Please upload .docx.")

    raise UnsupportedDocumentError("Unsupported file format. Please upload a PDF or DOCX file.")


def _extract_pdf_text(
    content: bytes,
    pdf_reader_factory: Callable[[BytesIO], object] | None = None,
) -> str:
    try:
        reader = (pdf_reader_factory or PdfReader)(BytesIO(content))
        parts = []
        for page in getattr(reader, "pages", []):
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
        return _normalize_text("\n\n".join(parts))
    except DocumentTextError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse PDF: {exc}") from exc


def _extract_docx_text(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append("\t".join(cells))

        return _normalize_text("\n\n".join(parts))
    except DocumentTextError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse DOCX: {exc}") from exc


def extract_document_text(
    filename: str,
    content_type: str | None,
    content: bytes,
    *,
    pdf_reader_factory: Callable[[BytesIO], object] | None = None,
) -> str:
    """Return normalized text extracted from a supported uploaded document."""
    if not content:
        raise EmptyDocumentTextError("Uploaded file is empty.")

    kind = _validate_type(filename, content_type)
    if kind == "pdf":
        text = _extract_pdf_text(content, pdf_reader_factory=pdf_reader_factory)
    else:
        text = _extract_docx_text(content)

    if not text.strip():
        raise EmptyDocumentTextError("No extractable text found in uploaded document.")
    return text
